import os
import uuid
from flask import Flask, render_template, request, send_from_directory, url_for, redirect
from docx import Document
import docx  # нужен для docx.shared.Pt и RGBColor

app = Flask(__name__)
app.config["GENERATED_DIR"] = os.path.join(os.path.dirname(__file__), "generated")
os.makedirs(app.config["GENERATED_DIR"], exist_ok=True)

# Можно расширять список — он просто помогает собрать значения из формы.
# Любые {Плейсхолдеры} из документов тоже заменятся, если придут во входном словаре.
VARIABLES = [
    "{ФИО}",
    "{Паспортные данные}",
    "{Адрес}",
    "{ИНН}",
    "{СНИЛС}",
    "{Банковские реквизиты}",
    "{id artist / id contract}",
    "{Электронная почта лицензиара}",
    "{Дата}",
    "{Процент вознаграждения}",
    "{Процент вознаграждения прописью}"
]

def replace_in_paragraph(paragraph, mapping):
    """Надёжная замена даже когда плейсхолдер порезан на несколько runs."""
    if not paragraph.runs:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    changed = False
    for k, v in mapping.items():
        if k in full_text:
            full_text = full_text.replace(k, v)
            changed = True
    if changed:
        # затираем ранны и пишем заново одним run (проще всего)
        for _ in range(len(paragraph.runs)):
            paragraph.runs[0].clear()  # очищаем текст и формат первого
            paragraph.runs[0].text = ""  # на всякий
            paragraph._p.remove(paragraph.runs[0]._r)  # удаляем run
        run = paragraph.add_run(full_text)
        run.font.name = 'Times New Roman'
        run.font.size = docx.shared.Pt(11)
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

def replace_in_table(table, mapping):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, mapping)

def replace_in_headers_footers(doc, mapping):
    for section in doc.sections:
        for hdr in [section.header, section.first_page_header, section.even_page_header]:
            if hdr:
                for p in hdr.paragraphs:
                    replace_in_paragraph(p, mapping)
                for t in hdr.tables:
                    replace_in_table(t, mapping)
        for ftr in [section.footer, section.first_page_footer, section.even_page_footer]:
            if ftr:
                for p in ftr.paragraphs:
                    replace_in_paragraph(p, mapping)
                for t in ftr.tables:
                    replace_in_table(t, mapping)

def replace_variables_in_docx(template_path, values_dict):
    doc = Document(template_path)
    # Поглощаем все пары вида {…}: если в values_dict нет значения — оставим как есть
    mapping = {k: v for k, v in values_dict.items() if v is not None and v != ""}

    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)

    for t in doc.tables:
        replace_in_table(t, mapping)

    replace_in_headers_footers(doc, mapping)
    return doc

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # Собираем значения из формы (поля должны называться без фигурных скобок)
        values = {}
        for var in VARIABLES:
            key = var.strip("{}")
            values[var] = request.form.get(key, "").strip()

        # Генерим уникальные имена, чтобы ссылки на скачивание работали параллельно
        token = uuid.uuid4().hex
        contract_name = f"Лицензионный договор_{token}.docx"
        appendix_name = f"Приложение к договору_{token}.docx"

        contract_doc = replace_variables_in_docx("contract_template.docx", values)
        appendix_doc = replace_variables_in_docx("appendix_template.docx", values)

        contract_path = os.path.join(app.config["GENERATED_DIR"], contract_name)
        appendix_path = os.path.join(app.config["GENERATED_DIR"], appendix_name)

        contract_doc.save(contract_path)
        appendix_doc.save(appendix_path)

        return render_template(
            "success.html",
            contract=contract_name,
            appendix=appendix_name
        )

    # Покажем форму
    return render_template("form.html", variables=[v.strip("{}") for v in VARIABLES])

@app.route("/download/<path:filename>")
def download(filename):
    return send_from_directory(app.config["GENERATED_DIR"], filename, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
